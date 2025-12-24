"""
Trajanus USA - Convert Markdown to Word DOCX
Usage: python convert_md_to_docx.py [file1.md] [file2.md] ...
If no files specified, converts all .md files in current directory

Requires: pip install python-docx markdown
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import os
import sys
import glob
import re

def md_to_docx(filepath):
    """Convert a single MD file to DOCX"""
    filename = os.path.basename(filepath)
    docx_name = filename.replace('.md', '.docx')
    
    print(f"  Converting: {filename}")
    
    # Read markdown content
    with open(filepath, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Process markdown line by line
    lines = md_content.split('\n')
    
    for line in lines:
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(text, style='List Number')
        # Code blocks (simplified)
        elif line.startswith('```'):
            continue  # Skip code fence markers
        # Horizontal rules
        elif line.startswith('---') or line.startswith('==='):
            doc.add_paragraph('_' * 50)
        # Regular paragraphs
        elif line.strip():
            # Remove basic markdown formatting
            clean = line.replace('**', '').replace('*', '').replace('`', '')
            doc.add_paragraph(clean)
    
    # Save
    doc.save(docx_name)
    print(f"    ✓ Created: {docx_name}")
    return docx_name

def main():
    print("\n" + "="*60)
    print("TRAJANUS - CONVERT MD TO WORD DOCX")
    print("="*60 + "\n")
    
    # Get files to convert
    if len(sys.argv) > 1:
        files = sys.argv[1:]
    else:
        files = glob.glob('*.md')
    
    if not files:
        print("No .md files found!")
        return
    
    print(f"Files to convert: {len(files)}\n")
    
    converted = []
    for filepath in files:
        if os.path.exists(filepath):
            docx = md_to_docx(filepath)
            converted.append(docx)
        else:
            print(f"  ⚠ File not found: {filepath}")
    
    print("\n" + "="*60)
    print(f"COMPLETE - Created {len(converted)} DOCX files")
    print("="*60 + "\n")

if __name__ == "__main__":
    try:
        main()
    except ImportError:
        print("\n❌ ERROR: python-docx not installed")
        print("Run: pip install python-docx --break-system-packages")
    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
