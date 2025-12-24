# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
"""
CONVERT_TO_DOCX.py
==================
Converts session files (.md, .txt) to beautiful .docx format.

Usage:
    python CONVERT_TO_DOCX.py diary_file.md
    python CONVERT_TO_DOCX.py technical_journal.md --type technical
    python CONVERT_TO_DOCX.py *.md

Author: Claude (for Bill King, Trajanus USA)
Date: December 6, 2025
"""

import os
import sys
import argparse
import glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CONFIGURATION
# ============================================================================
CONFIG = {
    # Output folder for .docx files (same location as input if not specified)
    'output_folder': None,  # Set to a path, or None to use same folder as input
    
    # Font settings
    'font': 'Architext',
    
    # Colors (RGB values)
    'colors': {
        'heading': (30, 58, 95),      # Navy blue
        'accent': (247, 148, 29),     # Orange
        'body': (44, 62, 80),         # Dark gray
    }
}


def convert_to_diary_docx(input_path, output_path=None):
    """Convert a file to beautiful diary-style .docx with Architext font."""
    
    # Read input file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Determine output path
    if output_path is None:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.docx'
    
    # Create document
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Daily Project Journal")
    run.font.name = CONFIG['font']
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Bill's Construction Management Log | Trajanus USA")
    run.font.name = CONFIG['font']
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*CONFIG['colors']['body'])
    
    # Separator
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("_" * 50)
    run.font.name = CONFIG['font']
    run.font.color.rgb = RGBColor(*CONFIG['colors']['accent'])
    
    # Date
    today = datetime.now().strftime("%A, %B %d, %Y")
    date_para = doc.add_paragraph()
    run = date_para.add_run(today)
    run.font.name = CONFIG['font']
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(*CONFIG['colors']['accent'])
    
    doc.add_paragraph()
    
    # Content - preserve paragraphs
    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.4)
        run = p.add_run(para_text)
        run.font.name = CONFIG['font']
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*CONFIG['colors']['body'])
    
    # Footer
    doc.add_paragraph()
    footer_sep = doc.add_paragraph()
    footer_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_sep.add_run("_" * 50)
    run.font.name = CONFIG['font']
    run.font.color.rgb = RGBColor(*CONFIG['colors']['accent'])
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Trajanus USA")
    run.font.name = CONFIG['font']
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
    
    # Save
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")
    return output_path


def convert_to_technical_docx(input_path, output_path=None):
    """Convert a file to technical journal style .docx."""
    
    # Read input file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Determine output path
    if output_path is None:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.docx'
    
    # Create document
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("TECHNICAL JOURNAL")
    run.font.name = 'Georgia'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
    
    # Subtitle
    today = datetime.now().strftime("%B %d, %Y")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Trajanus USA | {today}")
    run.font.name = 'Georgia'
    run.font.size = Pt(10)
    run.font.italic = True
    
    # Separator
    doc.add_paragraph("=" * 60)
    
    # Content
    lines = content.split('\n')
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        
        p = doc.add_paragraph()
        
        # Detect headers
        if line.startswith('#') or line.isupper():
            run = p.add_run(line.lstrip('#').strip())
            run.font.name = 'Georgia'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        # Detect code
        elif line.startswith('    ') or line.startswith('```'):
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        # Regular text
        else:
            run = p.add_run(line)
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
    
    # Save
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert session files to beautiful .docx format'
    )
    parser.add_argument('files', nargs='+', help='Files to convert (supports wildcards)')
    parser.add_argument('--type', '-t', choices=['diary', 'technical'], 
                        default='diary', help='Document style (default: diary)')
    parser.add_argument('--output', '-o', help='Output folder (default: same as input)')
    
    args = parser.parse_args()
    
    # Expand wildcards
    all_files = []
    for pattern in args.files:
        matches = glob.glob(pattern)
        if matches:
            all_files.extend(matches)
        else:
            all_files.append(pattern)  # Maybe it's an exact path
    
    if not all_files:
        print("[ERROR] No files found")
        return 1
    
    print(f"Converting {len(all_files)} file(s) to .docx...")
    print(f"Style: {args.type}")
    print(f"Font: {CONFIG['font']}")
    print("-" * 50)
    
    for filepath in all_files:
        if not os.path.exists(filepath):
            print(f"[SKIP] File not found: {filepath}")
            continue
        
        # Determine output path
        if args.output:
            filename = os.path.basename(filepath)
            base = os.path.splitext(filename)[0]
            output_path = os.path.join(args.output, base + '.docx')
        else:
            output_path = None  # Same folder as input
        
        # Convert based on type
        if args.type == 'diary':
            convert_to_diary_docx(filepath, output_path)
        else:
            convert_to_technical_docx(filepath, output_path)
    
    print("-" * 50)
    print("[DONE] Conversion complete!")
    return 0


if __name__ == "__main__":
    sys.exit(main())
