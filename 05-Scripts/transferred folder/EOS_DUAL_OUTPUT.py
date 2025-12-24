# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
"""
EOS_DUAL_OUTPUT.py
==================
END OF SESSION - DUAL OUTPUT EDITION

Creates TWO versions of each entry:
1. .DOCX file - Beautiful formatting with custom fonts (for Bill)
2. Google Doc append - Plain text for Claude's continuity

Same content. Two formats. One button.

Requirements:
    pip install python-docx google-auth google-auth-oauthlib google-api-python-client

Author: Claude (for Bill King, Trajanus USA)
Date: December 6, 2025
"""

import os
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

# Google API Scopes
SCOPES = [
    'https://www.googleapis.com/auth/documents',
    'https://www.googleapis.com/auth/drive'
]

# ============================================================================
# CONFIGURATION
# ============================================================================
CONFIG = {
    'credentials_path': r'G:\My Drive\00 - Trajanus USA\00-Command-Center\Credentials\credentials.json',
    'token_path': r'G:\My Drive\00 - Trajanus USA\00-Command-Center\token.json',
    
    # Where to save .docx files
    'docx_output_folder': r'G:\My Drive\00 - Trajanus USA\07-Session-Journals',
    
    # Custom fonts for .docx (change these to match your installed fonts)
    'fonts': {
        'heading': 'Architext',              # Bill's handwriting font for headers
        'body': 'Architext',                 # Bill's handwriting font for body too
        'fallback_heading': 'Comic Sans MS', # Fallback if font not installed
        'fallback_body': 'Courier New',      # Fallback if font not installed
    },
    
    # Colors (RGB values)
    'colors': {
        'heading': (30, 58, 95),      # Navy blue #1E3A5F
        'accent': (247, 148, 29),     # Orange #F7941D
        'body': (44, 62, 80),         # Dark gray #2C3E50
    }
}

# Google Doc IDs for appending
LIVING_DOCS = {
    'diary': {
        'id': '1HKOisNN8A5rf9YdFJnJSdgH326bdJTun2rDqObNvrM8',
        'name': 'Bills_Daily_Diary_MASTER'
    },
    'technical': {
        'id': '1iPZAmi2bYBRmDnsgwZK3UZFCsB_YHj9RvRtKWJqDb2Q',
        'name': 'Technical_Journal_MASTER'
    },
    'code': {
        'id': None,  # To be created
        'name': 'Code_Repository_MASTER'
    }
}


class DualOutputEOS:
    """End of Session with dual output: .docx and Google Docs."""
    
    def __init__(self):
        self.creds = None
        self.docs_service = None
        self.today = datetime.now().strftime("%Y-%m-%d")
        self.day_name = datetime.now().strftime("%A")
        self.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.authenticate()
    
    def authenticate(self):
        """Handle OAuth authentication for Google Docs."""
        token_path = CONFIG['token_path']
        creds_path = CONFIG['credentials_path']
        
        if os.path.exists(token_path):
            self.creds = Credentials.from_authorized_user_file(token_path, SCOPES)
        
        if not self.creds or not self.creds.valid:
            if self.creds and self.creds.expired and self.creds.refresh_token:
                self.creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(creds_path, SCOPES)
                self.creds = flow.run_local_server(port=0)
            
            with open(token_path, 'w') as token:
                token.write(self.creds.to_json())
        
        self.docs_service = build('docs', 'v1', credentials=self.creds)
        print("[AUTH] Connected to Google Docs API")
    
    # =========================================================================
    # DOCX CREATION (Beautiful formatted version for Bill)
    # =========================================================================
    
    def create_diary_docx(self, content, title=None):
        """Create a beautifully formatted .docx diary entry."""
        doc = Document()
        
        # Set up styles
        self._setup_diary_styles(doc)
        
        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("Daily Project Journal")
        run.font.name = CONFIG['fonts']['heading']
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Bill's Construction Management Log | Trajanus USA")
        run.font.name = CONFIG['fonts']['body']
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*CONFIG['colors']['body'])
        
        # Horizontal line
        doc.add_paragraph("_" * 70)
        
        # Date header
        date_para = doc.add_paragraph()
        run = date_para.add_run(f"{self.day_name}, {self.today}")
        run.font.name = CONFIG['fonts']['heading']
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(*CONFIG['colors']['accent'])
        
        # Title if provided
        if title:
            title_para = doc.add_paragraph()
            run = title_para.add_run(title)
            run.font.name = CONFIG['fonts']['heading']
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        
        # Separator
        doc.add_paragraph("-" * 50)
        
        # Content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.5)
                run = p.add_run(paragraph.strip())
                run.font.name = CONFIG['fonts']['body']
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*CONFIG['colors']['body'])
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("_" * 70)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Trajanus USA - Construction Project Management Excellence")
        run.font.name = CONFIG['fonts']['body']
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        
        # Save
        filename = f"Bills_Daily_Diary_{self.today}.docx"
        filepath = os.path.join(CONFIG['docx_output_folder'], 'Personal-Diaries', filename)
        
        # Ensure folder exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        doc.save(filepath)
        print(f"[DOCX] Created: {filename}")
        return filepath
    
    def create_technical_docx(self, content, title=None):
        """Create a Scientific American style technical journal entry."""
        doc = Document()
        
        # Header - Scientific American style
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("TECHNICAL JOURNAL")
        run.font.name = 'Georgia'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        
        # Volume/Issue style subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Trajanus USA | {self.today} | {self.day_name}")
        run.font.name = 'Georgia'
        run.font.size = Pt(10)
        run.font.italic = True
        
        # Double line separator
        doc.add_paragraph("=" * 70)
        
        # Entry title
        if title:
            title_para = doc.add_paragraph()
            run = title_para.add_run(title.upper())
            run.font.name = 'Georgia'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
        
        # Timestamp
        time_para = doc.add_paragraph()
        run = time_para.add_run(f"Recorded: {self.timestamp}")
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.italic = True
        
        doc.add_paragraph("-" * 50)
        
        # Content - detect sections and format accordingly
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect headers (lines in ALL CAPS or starting with ##)
            if line.isupper() or line.startswith('##'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('##', '').strip())
                run.font.name = 'Georgia'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(*CONFIG['colors']['heading'])
            
            # Detect code blocks (lines starting with spaces or containing code patterns)
            elif line.startswith('    ') or line.startswith('```'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("=" * 70)
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Confidential Technical Documentation | Trajanus USA")
        run.font.name = 'Georgia'
        run.font.size = Pt(8)
        run.font.italic = True
        
        # Save
        filename = f"Technical_Journal_{self.today}.docx"
        filepath = os.path.join(CONFIG['docx_output_folder'], 'Technical-Journals', filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        doc.save(filepath)
        print(f"[DOCX] Created: {filename}")
        return filepath
    
    def _setup_diary_styles(self, doc):
        """Set up custom styles for diary document."""
        # This could be expanded for more sophisticated styling
        pass
    
    # =========================================================================
    # GOOGLE DOC APPEND (Plain text for Claude)
    # =========================================================================
    
    def get_document_end_index(self, doc_id):
        """Get the index position at the end of a document."""
        try:
            doc = self.docs_service.documents().get(documentId=doc_id).execute()
            content = doc.get('body', {}).get('content', [])
            if content:
                last_element = content[-1]
                return last_element.get('endIndex', 1) - 1
            return 1
        except HttpError as e:
            print(f"[ERROR] Could not get document: {e}")
            return None
    
    def append_to_google_doc(self, doc_key, content, title=None):
        """Append plain text to Google Doc for Claude's continuity."""
        doc_info = LIVING_DOCS[doc_key]
        doc_id = doc_info['id']
        
        if not doc_id:
            print(f"[SKIP] {doc_info['name']} - no document ID configured")
            return False
        
        # Build plain text entry
        entry_parts = [
            "\n\n" + "=" * 70 + "\n",
            f"SESSION: {self.timestamp} ({self.day_name})\n"
        ]
        
        if title:
            entry_parts.append(f"TOPIC: {title}\n")
        
        entry_parts.append("-" * 70 + "\n\n")
        entry_parts.append(content)
        
        if not content.endswith("\n"):
            entry_parts.append("\n")
        
        full_entry = "".join(entry_parts)
        
        # Get end position
        end_index = self.get_document_end_index(doc_id)
        if end_index is None:
            return False
        
        # Insert
        requests = [{
            'insertText': {
                'location': {'index': end_index},
                'text': full_entry
            }
        }]
        
        try:
            self.docs_service.documents().batchUpdate(
                documentId=doc_id,
                body={'requests': requests}
            ).execute()
            
            print(f"[GDOC] Appended to {doc_info['name']}")
            return True
            
        except HttpError as e:
            print(f"[ERROR] Failed: {e}")
            return False
    
    # =========================================================================
    # MAIN WORKFLOW
    # =========================================================================
    
    def get_multiline_input(self, prompt):
        """Get multi-line input from user."""
        print(prompt)
        print("(Type or paste your content. Press Enter twice to finish.)")
        print("-" * 50)
        
        lines = []
        empty_count = 0
        
        while True:
            try:
                line = input()
                if line == "":
                    empty_count += 1
                    if empty_count >= 2:
                        break
                    lines.append("")
                else:
                    empty_count = 0
                    lines.append(line)
            except EOFError:
                break
        
        return "\n".join(lines).strip()
    
    def run_eos(self, diary_content=None, tech_content=None, diary_title=None, tech_title=None):
        """Run the full End of Session process."""
        print("=" * 70)
        print("TRAJANUS END OF SESSION - DUAL OUTPUT")
        print("=" * 70)
        print(f"Date: {self.today} ({self.day_name})")
        print(f"Time: {datetime.now().strftime('%H:%M')}")
        print()
        print("This will create:")
        print("  1. Beautiful .DOCX files (with your custom fonts)")
        print("  2. Google Doc entries (for Claude's continuity)")
        print("=" * 70)
        
        results = {
            'diary_docx': None,
            'diary_gdoc': False,
            'tech_docx': None,
            'tech_gdoc': False
        }
        
        # === DIARY ===
        print("\n" + "=" * 70)
        print("SECTION 1: DAILY DIARY")
        print("=" * 70)
        
        if diary_content is None:
            print("\nWhat happened today? How do you feel about the session?")
            print("(This is YOUR voice - emotions, reflections, the human story)")
            diary_content = self.get_multiline_input("")
        
        if diary_content:
            # Create both versions
            results['diary_docx'] = self.create_diary_docx(diary_content, diary_title)
            results['diary_gdoc'] = self.append_to_google_doc('diary', diary_content, diary_title)
        else:
            print("[SKIP] No diary entry provided")
        
        # === TECHNICAL JOURNAL ===
        print("\n" + "=" * 70)
        print("SECTION 2: TECHNICAL JOURNAL")
        print("=" * 70)
        
        if tech_content is None:
            print("\nWhat technical work was done? Decisions made? Problems solved?")
            tech_content = self.get_multiline_input("")
        
        if tech_content:
            # Create both versions
            results['tech_docx'] = self.create_technical_docx(tech_content, tech_title)
            results['tech_gdoc'] = self.append_to_google_doc('technical', tech_content, tech_title)
        else:
            print("[SKIP] No technical entry provided")
        
        # === SUMMARY ===
        print("\n" + "=" * 70)
        print("EOS COMPLETE - DUAL OUTPUT SUMMARY")
        print("=" * 70)
        
        if results['diary_docx']:
            print(f"  [DOCX] Diary: {os.path.basename(results['diary_docx'])}")
        if results['diary_gdoc']:
            print(f"  [GDOC] Diary: Appended to master")
        if results['tech_docx']:
            print(f"  [DOCX] Technical: {os.path.basename(results['tech_docx'])}")
        if results['tech_gdoc']:
            print(f"  [GDOC] Technical: Appended to master")
        
        print()
        print("[SUCCESS] Your beautiful .docx files are ready!")
        print("[SUCCESS] Claude will have full context next session!")
        print("=" * 70)
        
        return results


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='End of Session - Dual Output (DOCX + Google Docs)'
    )
    parser.add_argument('--diary', '-d', type=str, help='Diary content or file path')
    parser.add_argument('--tech', '-t', type=str, help='Technical journal content or file path')
    parser.add_argument('--diary-title', type=str, help='Title for diary entry')
    parser.add_argument('--tech-title', type=str, help='Title for technical entry')
    
    args = parser.parse_args()
    
    # Handle file inputs
    diary_content = None
    tech_content = None
    
    if args.diary:
        if os.path.exists(args.diary):
            with open(args.diary, 'r', encoding='utf-8') as f:
                diary_content = f.read()
        else:
            diary_content = args.diary
    
    if args.tech:
        if os.path.exists(args.tech):
            with open(args.tech, 'r', encoding='utf-8') as f:
                tech_content = f.read()
        else:
            tech_content = args.tech
    
    # Run EOS
    eos = DualOutputEOS()
    eos.run_eos(
        diary_content=diary_content,
        tech_content=tech_content,
        diary_title=args.diary_title,
        tech_title=args.tech_title
    )


if __name__ == "__main__":
    main()
