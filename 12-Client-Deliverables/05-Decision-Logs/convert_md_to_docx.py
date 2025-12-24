#!/usr/bin/env python3
"""
MARKDOWN TO DOCX CONVERTER
Converts markdown files to professionally formatted Word documents with Trajanus branding.

Author: Claude Code
Date: 2025-12-18
"""

import os
import re
import pickle
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# ============================================================================
# CONFIGURATION
# ============================================================================

CREDENTIALS_PATH = Path("G:/My Drive/00 - Trajanus USA/00-Command-Center/Credentials/token.pickle")
OUTPUT_DIR = Path("G:/My Drive/00 - Trajanus USA/00-Command-Center/outputs")
LIVING_DOCS_DIR = Path("G:/My Drive/00 - Trajanus USA/03-Living-Documents")

# Trajanus Colors (RGB)
TRAJANUS_BROWN = RGBColor(155, 126, 82)
TRAJANUS_BROWN_DARK = RGBColor(123, 97, 66)
TEXT_PRIMARY = RGBColor(51, 51, 51)
TEXT_SECONDARY = RGBColor(102, 102, 102)

# Google Drive folder ID for uploads (03-Living-Documents)
LIVING_DOCS_FOLDER_ID = "1R2Fm9vA52_YJF3beKBCZsa3Cp0F7CLTt"

# ============================================================================
# CONVERSION FUNCTIONS
# ============================================================================

def convert_md_to_docx(md_file, output_file):
    """Convert markdown to formatted Word document with Trajanus branding."""

    doc = Document()

    # Set document properties
    title = Path(md_file).stem.replace('_', ' ').replace('-', ' ')
    doc.core_properties.title = title
    doc.core_properties.author = 'Trajanus USA'
    doc.core_properties.company = 'Trajanus USA'

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        line_stripped = line.rstrip()

        # Handle code blocks
        if line_stripped.startswith('```'):
            if in_code_block:
                # End of code block - add accumulated code
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = TEXT_PRIMARY
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip empty lines (add paragraph break)
        if not line_stripped:
            doc.add_paragraph()
            continue

        # H1 headers (# )
        if line_stripped.startswith('# '):
            text = line_stripped[2:]
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                run.font.color.rgb = TRAJANUS_BROWN
                run.font.size = Pt(18)

        # H2 headers (## )
        elif line_stripped.startswith('## '):
            text = line_stripped[3:]
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                run.font.color.rgb = TRAJANUS_BROWN
                run.font.size = Pt(14)

        # H3 headers (### )
        elif line_stripped.startswith('### '):
            text = line_stripped[4:]
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.color.rgb = TRAJANUS_BROWN_DARK
                run.font.size = Pt(12)

        # H4 headers (#### )
        elif line_stripped.startswith('#### '):
            text = line_stripped[5:]
            h = doc.add_heading(text, level=4)
            for run in h.runs:
                run.font.color.rgb = TRAJANUS_BROWN_DARK
                run.font.size = Pt(11)
                run.font.bold = True

        # Bullet lists (- or *)
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            text = line_stripped[2:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)

        # Nested bullet lists (  - or  *)
        elif line_stripped.startswith('  - ') or line_stripped.startswith('  * '):
            text = line_stripped[4:]
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet 2')
            add_formatted_text(p, text)

        # Numbered lists (1. )
        elif re.match(r'^\d+\.\s', line_stripped):
            text = re.sub(r'^\d+\.\s', '', line_stripped)
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Horizontal rule (--- or ***)
        elif line_stripped in ['---', '***', '___']:
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = TRAJANUS_BROWN

        # Table rows (|...|)
        elif line_stripped.startswith('|') and line_stripped.endswith('|'):
            # Simple table handling - just add as formatted paragraph
            text = line_stripped.replace('|', ' | ').strip()
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)

        # Regular paragraph
        else:
            text = process_inline_formatting(line_stripped)
            p = doc.add_paragraph()
            add_formatted_text(p, text)

    # Save document
    doc.save(output_file)
    return True


def process_inline_formatting(text):
    """Process inline markdown formatting markers."""
    return text


def add_formatted_text(paragraph, text):
    """Add text with inline formatting to paragraph."""

    # Handle bold (**text**) and italic (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_PRIMARY
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_PRIMARY
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_SECONDARY
        else:
            # Regular text
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_PRIMARY


# ============================================================================
# GOOGLE DRIVE FUNCTIONS
# ============================================================================

def get_drive_service():
    """Get Google Drive service."""
    with open(CREDENTIALS_PATH, "rb") as token:
        creds = pickle.load(token)
    return build("drive", "v3", credentials=creds)


def upload_to_drive(service, file_path, folder_id):
    """Upload file to Google Drive and convert to Google Docs."""

    file_name = Path(file_path).name

    # Check if file already exists in folder
    query = f"'{folder_id}' in parents and name = '{file_name}' and trashed = false"
    results = service.files().list(q=query, fields="files(id, name)").execute()
    existing = results.get('files', [])

    if existing:
        # Delete existing file
        service.files().delete(fileId=existing[0]['id']).execute()
        print(f"  [REPLACED] Deleted existing: {file_name}")

    # Upload new file
    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }

    media = MediaFileUpload(
        file_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=True
    )

    file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id, name, webViewLink'
    ).execute()

    return file


# ============================================================================
# MAIN CONVERSION
# ============================================================================

def find_md_files():
    """Find all markdown files to convert."""

    md_files = []

    # Priority 1: Code_Repository_2025-12-18.md (today's file)
    today_file = LIVING_DOCS_DIR / "Code_Repository_2025-12-18.md"
    if today_file.exists():
        md_files.append(str(today_file))

    # Priority 2: Files in "md docs to convert" folder
    convert_folder = LIVING_DOCS_DIR / "01 Code-Repository" / "Code_Repository_Daily_Entries" / "md docs to convert"
    if convert_folder.exists():
        for f in convert_folder.glob("*.md"):
            md_files.append(str(f))

    return md_files


def main():
    print("=" * 70)
    print("MARKDOWN TO DOCX CONVERTER")
    print("=" * 70)
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print()

    # Create output directory
    docx_output_dir = OUTPUT_DIR / "docx_conversions"
    docx_output_dir.mkdir(parents=True, exist_ok=True)

    # Find files to convert
    md_files = find_md_files()

    if not md_files:
        print("No markdown files found to convert.")
        return

    print(f"Found {len(md_files)} markdown file(s) to convert")
    print()

    # Convert files
    print("[PHASE 1] CONVERTING TO DOCX")
    print("-" * 50)

    converted = []
    errors = []

    for md_file in md_files:
        file_name = Path(md_file).stem
        output_file = docx_output_dir / f"{file_name}.docx"

        try:
            print(f"  Converting: {Path(md_file).name}")
            convert_md_to_docx(md_file, str(output_file))
            converted.append({
                'source': md_file,
                'output': str(output_file),
                'name': file_name
            })
            print(f"    [OK] -> {output_file.name}")
        except Exception as e:
            errors.append({
                'file': md_file,
                'error': str(e)
            })
            print(f"    [ERROR] {e}")

    print()
    print(f"Converted: {len(converted)} | Errors: {len(errors)}")

    # Upload to Google Drive
    if converted:
        print()
        print("[PHASE 2] UPLOADING TO GOOGLE DRIVE")
        print("-" * 50)

        service = get_drive_service()
        uploaded = []

        for item in converted:
            try:
                print(f"  Uploading: {Path(item['output']).name}")
                result = upload_to_drive(service, item['output'], LIVING_DOCS_FOLDER_ID)
                uploaded.append({
                    'name': result['name'],
                    'id': result['id'],
                    'link': result.get('webViewLink', 'N/A')
                })
                print(f"    [OK] Uploaded to Google Drive")
            except Exception as e:
                print(f"    [ERROR] {e}")

        print()
        print(f"Uploaded: {len(uploaded)} file(s)")

    # Generate report
    print()
    print("=" * 70)
    print("CONVERSION REPORT")
    print("=" * 70)
    print()

    print("CONVERTED FILES:")
    for item in converted:
        size = os.path.getsize(item['output'])
        print(f"  [OK] {Path(item['output']).name} ({size:,} bytes)")

    if errors:
        print()
        print("ERRORS:")
        for err in errors:
            print(f"  [ERROR] {Path(err['file']).name}: {err['error']}")

    print()
    print(f"Total converted: {len(converted)}")
    print(f"Total errors: {len(errors)}")
    print(f"Output directory: {docx_output_dir}")
    print()
    print("Next steps:")
    print("1. Bill runs CONVERT_NEW_FILES_ONLY.ps1 to create Google Docs versions")
    print("2. Verify Google Docs versions created")
    print("3. Archive original .md files if desired")


if __name__ == "__main__":
    main()
